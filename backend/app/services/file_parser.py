"""File parser service for extracting text from various document formats."""
import io
from typing import Optional


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from a Word .docx file."""
    import docx
    doc = docx.Document(io.BytesIO(file_bytes))
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text)
    
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                if cell.text.strip():
                    row_text.append(cell.text.strip())
            if row_text:
                paragraphs.append(" | ".join(row_text))
    
    return "\n".join(paragraphs)


def extract_text_from_xlsx(file_bytes: bytes) -> str:
    """Extract text from an Excel .xlsx file."""
    import openpyxl
    wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
    all_text = []
    
    for sheet in wb.worksheets:
        all_text.append(f"=== {sheet.title} ===")
        for row in sheet.iter_rows(values_only=True):
            cells = [str(cell) for cell in row if cell is not None]
            if cells:
                all_text.append(" | ".join(cells))
    
    return "\n".join(all_text)


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF file."""
    import pymupdf
    doc = pymupdf.open(stream=file_bytes, filetype="pdf")
    text_parts = []
    
    for page_num, page in enumerate(doc):
        page_text = page.get_text()
        if page_text.strip():
            text_parts.append(f"--- Page {page_num + 1} ---\n{page_text}")
    
    doc.close()
    return "\n".join(text_parts)


def extract_text(file_bytes: bytes, file_type: str) -> Optional[str]:
    """
    Extract text from a file based on its type.
    
    Args:
        file_bytes: The raw file bytes
        file_type: File extension (docx, xlsx, pdf)
    
    Returns:
        Extracted text or None if format not supported
    """
    file_type = file_type.lower().strip(".")
    
    extractors = {
        "docx": extract_text_from_docx,
        "xlsx": extract_text_from_xlsx,
        "xls": extract_text_from_xlsx,
        "pdf": extract_text_from_pdf,
    }
    
    extractor = extractors.get(file_type)
    if extractor:
        return extractor(file_bytes)
    
    # Try to read as plain text
    try:
        return file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        return None
